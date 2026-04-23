#!/usr/bin/env python3
"""Convert resume .docx files in docs/resume/ to Markdown in resume-data/source-docx-md/."""

import re
import sys
from pathlib import Path

from docx import Document


def style_to_md(text: str, style_name: str) -> str:
    s = (style_name or "").lower()
    text = text.rstrip()
    if not text:
        return ""
    if "heading 1" in s or s == "title":
        return f"# {text}"
    if "heading 2" in s:
        return f"## {text}"
    if "heading 3" in s:
        return f"### {text}"
    if "heading 4" in s:
        return f"#### {text}"
    if "list bullet" in s or "list paragraph" in s:
        return f"- {text}"
    return text


def runs_to_md(paragraph) -> str:
    out = []
    for run in paragraph.runs:
        t = run.text
        if not t:
            continue
        if run.bold and run.italic:
            out.append(f"***{t}***")
        elif run.bold:
            out.append(f"**{t}**")
        elif run.italic:
            out.append(f"*{t}*")
        else:
            out.append(t)
    return "".join(out)


def convert(docx_path: Path) -> str:
    doc = Document(docx_path)
    lines: list[str] = []
    for p in doc.paragraphs:
        raw = runs_to_md(p) or p.text
        md = style_to_md(raw, p.style.name if p.style else "")
        if md:
            lines.append(md)
        elif lines and lines[-1] != "":
            lines.append("")

    for table in doc.tables:
        lines.append("")
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
        lines.append("")

    out = "\n".join(lines)
    out = re.sub(r"\n{3,}", "\n\n", out)
    return out.strip() + "\n"


def main() -> int:
    root = Path(__file__).resolve().parent.parent
    src = root / "docs" / "resume"
    dst = root / "resume-data" / "source-docx-md"
    dst.mkdir(parents=True, exist_ok=True)

    files = sorted(src.glob("*.docx"))
    if not files:
        print(f"No .docx files in {src}", file=sys.stderr)
        return 1

    for f in files:
        md = convert(f)
        out_name = re.sub(r"[^A-Za-z0-9._-]+", "_", f.stem) + ".md"
        out_path = dst / out_name
        out_path.write_text(md, encoding="utf-8")
        print(f"  {f.name} -> {out_path.relative_to(root)}  ({len(md)} chars)")

    print(f"\nConverted {len(files)} files to {dst.relative_to(root)}/")
    return 0


if __name__ == "__main__":
    sys.exit(main())

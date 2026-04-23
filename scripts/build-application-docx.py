#!/usr/bin/env python3
"""Build .docx artifacts for a job application from Markdown sources.

For each application slug, produces:
  - tailored-resumes/<slug>/resume.docx       (V2_McDonald's-styled)
  - tailored-resumes/<slug>/resume-ats.docx   (single-column ATS-safe)
  - cover-letters/<slug>.docx                 (single-column)

Usage:
  python3 scripts/build-application-docx.py <slug>

Example:
  python3 scripts/build-application-docx.py hpe__senior-pm-greenlake-flex__2026-04-22
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm

ROOT = Path(__file__).resolve().parent.parent

NAVY = RGBColor(0x1E, 0x3A, 0x5F)
GOLD = RGBColor(0xD4, 0xAF, 0x37)
SLATE = RGBColor(0x5A, 0x7A, 0x9F)
GRAY = RGBColor(0x6B, 0x72, 0x80)


def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color: str = "1E3A5F", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for border_name in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def configure_page(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)


def set_default_font(doc: Document, name: str = "Calibri", size: int = 10) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def add_run(paragraph, text: str, *, bold=False, italic=False, size: float | None = None,
            color: RGBColor | None = None, font: str | None = None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if font is not None:
        run.font.name = font
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), font)
        rfonts.set(qn("w:hAnsi"), font)
    return run


def section_heading(doc: Document, text: str, *, color: RGBColor = NAVY, size: int = 12) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text.upper(), bold=True, size=size, color=color)
    add_horizontal_rule(p, color="1E3A5F")


def add_horizontal_rule(paragraph, color: str = "1E3A5F") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def parse_resume_md(md_text: str) -> dict:
    """Parse the tailored resume markdown into structured data."""
    sections: dict[str, list[str]] = {}
    current = "header"
    sections[current] = []
    for line in md_text.splitlines():
        m = re.match(r"^##\s+(.+)$", line)
        if m and not line.startswith("###"):
            current = m.group(1).strip()
            sections[current] = []
            continue
        sections[current].append(line)
    return sections


def split_experience_blocks(lines: list[str]) -> list[dict]:
    """Group experience lines into role blocks."""
    blocks = []
    current: dict | None = None
    for line in lines:
        line = line.rstrip()
        if line.startswith("### "):
            if current:
                blocks.append(current)
            current = {"company": line[4:].strip(), "role_line": "", "bullets": []}
        elif line.startswith("**") and current is not None and not current["role_line"]:
            current["role_line"] = line.strip()
        elif line.startswith("- ") and current is not None:
            current["bullets"].append(line[2:].strip())
    if current:
        blocks.append(current)
    return blocks


def render_inline_bold(paragraph, text: str, *, size: float | None = None) -> None:
    """Render a text run that may contain **bold** markdown segments."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            add_run(paragraph, part[2:-2], bold=True, size=size)
        elif part:
            add_run(paragraph, part, size=size)


def write_styled_resume(slug: str) -> Path:
    """Build the V2_McDonald's-styled resume.docx."""
    src = ROOT / "tailored-resumes" / slug / "resume.md"
    dst = ROOT / "tailored-resumes" / slug / "resume.docx"
    md = src.read_text(encoding="utf-8")
    sections = parse_resume_md(md)

    doc = Document()
    configure_page(doc)
    set_default_font(doc, "Calibri", 10)

    # --- Header block ---
    header_lines = [l for l in sections.get("header", []) if l.strip()]
    name = next((l for l in header_lines if l.startswith("# ")), "# ABDOULAYE SOW")[2:]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, name.strip(), bold=True, size=22, color=NAVY)

    tagline_line = next((l for l in header_lines if l.startswith("**") and "|" in l), "")
    if tagline_line:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, tagline_line.strip("*").strip(), italic=True, size=10.5, color=SLATE)

    contact_lines = [l for l in header_lines if l and not l.startswith("#") and not l.startswith("**") and not l.startswith("---")]
    for cl in contact_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, cl.strip(), size=9.5, color=GRAY)

    # --- Professional Summary ---
    if "PROFESSIONAL SUMMARY" in sections:
        section_heading(doc, "Professional Summary")
        text = " ".join(l.strip() for l in sections["PROFESSIONAL SUMMARY"] if l.strip() and not l.startswith("---"))
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        render_inline_bold(p, text, size=10)

    # --- Key Achievements (2-cell table) ---
    if "KEY ACHIEVEMENTS" in sections:
        section_heading(doc, "Key Achievements")
        bullets = [l[2:].strip() for l in sections["KEY ACHIEVEMENTS"] if l.startswith("- ")]
        # Distribute into 2-column table
        mid = (len(bullets) + 1) // 2
        col1, col2 = bullets[:mid], bullets[mid:]
        rows = max(len(col1), len(col2))
        col1 += [""] * (rows - len(col1))
        col2 += [""] * (rows - len(col2))
        table = doc.add_table(rows=rows, cols=2)
        table.autofit = True
        for i in range(rows):
            for j, content in enumerate((col1[i], col2[i])):
                cell = table.rows[i].cells[j]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                set_cell_shading(cell, "F4F0E6")
                set_cell_borders(cell, color="D4AF37", size="4")
                cell_p = cell.paragraphs[0]
                cell_p.paragraph_format.space_before = Pt(2)
                cell_p.paragraph_format.space_after = Pt(2)
                if content:
                    add_run(cell_p, "▸  ", bold=True, color=GOLD, size=9.5)
                    render_inline_bold(cell_p, content, size=9.5)

    # --- Technical & Leadership Competencies (subsection table) ---
    if "TECHNICAL & LEADERSHIP COMPETENCIES" in sections:
        section_heading(doc, "Technical & Leadership Competencies")
        comp_lines = [l for l in sections["TECHNICAL & LEADERSHIP COMPETENCIES"] if l.strip() and not l.startswith("---")]
        i = 0
        rows_data = []
        while i < len(comp_lines):
            line = comp_lines[i].strip()
            if line.startswith("**") and line.endswith("**"):
                category = line[2:-2]
                content_parts = []
                j = i + 1
                while j < len(comp_lines) and not (comp_lines[j].strip().startswith("**") and comp_lines[j].strip().endswith("**")):
                    if comp_lines[j].strip():
                        content_parts.append(comp_lines[j].strip())
                    j += 1
                rows_data.append((category, " ".join(content_parts)))
                i = j
            else:
                i += 1
        if rows_data:
            table = doc.add_table(rows=len(rows_data), cols=2)
            table.autofit = True
            for r, (cat, content) in enumerate(rows_data):
                left = table.rows[r].cells[0]
                right = table.rows[r].cells[1]
                set_cell_shading(left, "1E3A5F")
                set_cell_borders(left, color="1E3A5F", size="4")
                set_cell_borders(right, color="1E3A5F", size="4")
                lp = left.paragraphs[0]
                lp.paragraph_format.space_before = Pt(2)
                lp.paragraph_format.space_after = Pt(2)
                add_run(lp, cat, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=9.5)
                rp = right.paragraphs[0]
                rp.paragraph_format.space_before = Pt(2)
                rp.paragraph_format.space_after = Pt(2)
                render_inline_bold(rp, content, size=9.5)
            # Set column widths
            for row in table.rows:
                row.cells[0].width = Cm(4.5)
                row.cells[1].width = Cm(13)

    # --- Professional Experience ---
    if "PROFESSIONAL EXPERIENCE" in sections:
        section_heading(doc, "Professional Experience")
        blocks = split_experience_blocks(sections["PROFESSIONAL EXPERIENCE"])
        for blk in blocks:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            add_run(p, blk["company"], bold=True, size=10.5, color=NAVY)

            if blk["role_line"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                # role_line looks like **Role** | *dates*
                role_line = blk["role_line"]
                role_match = re.match(r"\*\*(.+?)\*\*\s*\|\s*\*(.+?)\*", role_line)
                if role_match:
                    add_run(p, role_match.group(1), bold=True, italic=True, size=10, color=SLATE)
                    add_run(p, "  |  ", size=10, color=GRAY)
                    add_run(p, role_match.group(2), italic=True, size=10, color=GRAY)
                else:
                    add_run(p, role_line.replace("**", "").replace("*", ""), italic=True, size=10, color=SLATE)

            for bullet in blk["bullets"]:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.left_indent = Cm(0.5)
                render_inline_bold(p, bullet, size=10)

    # --- Certifications ---
    if "CERTIFICATIONS" in sections:
        section_heading(doc, "Certifications")
        for line in sections["CERTIFICATIONS"]:
            line = line.strip()
            if line.startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.left_indent = Cm(0.5)
                render_inline_bold(p, line[2:], size=10)

    # --- Education ---
    if "EDUCATION" in sections:
        section_heading(doc, "Education")
        for line in sections["EDUCATION"]:
            line = line.strip()
            if not line or line.startswith("---"):
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            render_inline_bold(p, line, size=10)

    # --- Publications ---
    if "PUBLICATIONS & THOUGHT LEADERSHIP" in sections:
        section_heading(doc, "Publications & Thought Leadership")
        for line in sections["PUBLICATIONS & THOUGHT LEADERSHIP"]:
            line = line.strip()
            if line.startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.left_indent = Cm(0.5)
                render_inline_bold(p, line[2:], size=9.5)

    # --- Additional Information ---
    if "ADDITIONAL INFORMATION" in sections:
        section_heading(doc, "Additional Information")
        for line in sections["ADDITIONAL INFORMATION"]:
            line = line.strip()
            if line.startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.left_indent = Cm(0.5)
                render_inline_bold(p, line[2:], size=9.5)

    doc.save(dst)
    return dst


def write_ats_resume(slug: str) -> Path:
    """Build a single-column ATS-safe resume.docx (no tables, plain styling).

    Prefers resume-ats.md as the source if present (lets the ATS version diverge
    in content from the visual resume.md — e.g., shorter tagline, fewer roles,
    combined Certs+Education section). Falls back to resume.md if not.
    """
    ats_src = ROOT / "tailored-resumes" / slug / "resume-ats.md"
    src = ats_src if ats_src.exists() else ROOT / "tailored-resumes" / slug / "resume.md"
    dst = ROOT / "tailored-resumes" / slug / "resume-ats.docx"
    md = src.read_text(encoding="utf-8")
    sections = parse_resume_md(md)

    doc = Document()
    configure_page(doc)
    set_default_font(doc, "Calibri", 11)

    # Header
    header_lines = [l for l in sections.get("header", []) if l.strip()]
    name = next((l for l in header_lines if l.startswith("# ")), "# ABDOULAYE SOW")[2:]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, name.strip(), bold=True, size=18)

    tagline_line = next((l for l in header_lines if l.startswith("**") and "|" in l), "")
    if tagline_line:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, tagline_line.strip("*").strip(), size=11)

    contact_lines = [l for l in header_lines if l and not l.startswith("#") and not l.startswith("**") and not l.startswith("---")]
    for cl in contact_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, cl.strip(), size=10)

    def heading(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, text.upper(), bold=True, size=12)
        add_horizontal_rule(p, color="000000")

    # Summary
    if "PROFESSIONAL SUMMARY" in sections:
        heading("Professional Summary")
        text = " ".join(l.strip() for l in sections["PROFESSIONAL SUMMARY"] if l.strip() and not l.startswith("---"))
        p = doc.add_paragraph()
        render_inline_bold(p, text, size=11)

    # Key Achievements (as bullet list, no table)
    if "KEY ACHIEVEMENTS" in sections:
        heading("Key Achievements")
        for line in sections["KEY ACHIEVEMENTS"]:
            if line.startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(0.5)
                render_inline_bold(p, line[2:], size=11)

    # Skills (Competencies, flattened to category: items)
    if "TECHNICAL & LEADERSHIP COMPETENCIES" in sections:
        heading("Skills")
        comp_lines = [l for l in sections["TECHNICAL & LEADERSHIP COMPETENCIES"] if l.strip() and not l.startswith("---")]
        i = 0
        while i < len(comp_lines):
            line = comp_lines[i].strip()
            if line.startswith("**") and line.endswith("**"):
                category = line[2:-2]
                content_parts = []
                j = i + 1
                while j < len(comp_lines) and not (comp_lines[j].strip().startswith("**") and comp_lines[j].strip().endswith("**")):
                    if comp_lines[j].strip():
                        content_parts.append(comp_lines[j].strip())
                    j += 1
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                add_run(p, f"{category}: ", bold=True, size=11)
                add_run(p, " ".join(content_parts), size=11)
                i = j
            else:
                i += 1

    # Experience
    if "PROFESSIONAL EXPERIENCE" in sections:
        heading("Professional Experience")
        blocks = split_experience_blocks(sections["PROFESSIONAL EXPERIENCE"])
        for blk in blocks:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            add_run(p, blk["company"], bold=True, size=11)

            if blk["role_line"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                role_line = blk["role_line"]
                role_match = re.match(r"\*\*(.+?)\*\*\s*\|\s*\*(.+?)\*", role_line)
                if role_match:
                    add_run(p, role_match.group(1), bold=True, size=11)
                    add_run(p, "  |  ", size=11)
                    add_run(p, role_match.group(2), italic=True, size=11)
                else:
                    add_run(p, role_line.replace("**", "").replace("*", ""), italic=True, size=11)

            for bullet in blk["bullets"]:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(0.5)
                render_inline_bold(p, bullet, size=11)

    # Certs (and possibly combined Certs & Education)
    combined_cert_edu = "CERTIFICATIONS & EDUCATION" in sections
    if combined_cert_edu:
        heading("Certifications and Education")
        for line in sections["CERTIFICATIONS & EDUCATION"]:
            line = line.strip()
            if line.startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(0.5)
                render_inline_bold(p, line[2:], size=11)
            elif line and not line.startswith("---"):
                p = doc.add_paragraph()
                render_inline_bold(p, line, size=11)
    else:
        if "CERTIFICATIONS" in sections:
            heading("Certifications")
            for line in sections["CERTIFICATIONS"]:
                line = line.strip()
                if line.startswith("- "):
                    p = doc.add_paragraph(style="List Bullet")
                    p.paragraph_format.left_indent = Cm(0.5)
                    render_inline_bold(p, line[2:], size=11)

        if "EDUCATION" in sections:
            heading("Education")
            for line in sections["EDUCATION"]:
                line = line.strip()
                if not line or line.startswith("---"):
                    continue
                p = doc.add_paragraph()
                render_inline_bold(p, line, size=11)

    # Publications
    if "PUBLICATIONS & THOUGHT LEADERSHIP" in sections:
        heading("Publications and Thought Leadership")
        for line in sections["PUBLICATIONS & THOUGHT LEADERSHIP"]:
            line = line.strip()
            if line.startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(0.5)
                render_inline_bold(p, line[2:], size=11)

    # Additional
    if "ADDITIONAL INFORMATION" in sections:
        heading("Additional Information")
        for line in sections["ADDITIONAL INFORMATION"]:
            line = line.strip()
            if line.startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(0.5)
                render_inline_bold(p, line[2:], size=11)

    doc.save(dst)
    return dst


def write_cover_letter(slug: str) -> Path:
    """Build the cover letter .docx (single column, classic letter format)."""
    src = ROOT / "cover-letters" / f"{slug}.md"
    dst = ROOT / "cover-letters" / f"{slug}.docx"
    md = src.read_text(encoding="utf-8")

    doc = Document()
    configure_page(doc)
    set_default_font(doc, "Calibri", 11)

    for raw_line in md.splitlines():
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph()
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        # Detect bold-only lines (header)
        if line.startswith("**") and line.endswith("**") and line.count("**") == 2:
            add_run(p, line[2:-2], bold=True, size=12)
        else:
            render_inline_bold(p, line, size=11)

    doc.save(dst)
    return dst


def main() -> int:
    if len(sys.argv) != 2:
        print(__doc__, file=sys.stderr)
        return 2
    slug = sys.argv[1]

    resume_styled = write_styled_resume(slug)
    print(f"  styled resume -> {resume_styled.relative_to(ROOT)}")

    resume_ats = write_ats_resume(slug)
    print(f"  ATS resume    -> {resume_ats.relative_to(ROOT)}")

    cover = write_cover_letter(slug)
    print(f"  cover letter  -> {cover.relative_to(ROOT)}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
